import logging
import os

from langchain_community.document_loaders import PDFPlumberLoader
from langchain.schema import Document
from docx import Document as DocxDocument

from src.document.splitter import split_documents

logger = logging.getLogger(__name__)

def process_document(file_path: str, chunk_size: int, chunk_overlap: int):
    """Đọc file PDF hoặc DOCX và trả về danh sách chunks"""

    ext = os.path.splitext(file_path)[1].lower()

    # ====================== LOAD FILE ======================
    if ext == ".pdf":
        loader = PDFPlumberLoader(file_path)
        raw_docs = loader.load()

    elif ext == ".docx":
        doc = DocxDocument(file_path)
        full_text = "\n".join([para.text for para in doc.paragraphs])
        raw_docs = [Document(page_content=full_text)]

    else:
        raise ValueError("Chỉ hỗ trợ PDF và DOCX")

    # ====================== SPLIT (gọi từ splitter.py) ======================
    documents = split_documents(raw_docs, chunk_size, chunk_overlap)

    # Không dùng logger theo yêu cầu đề
    return documents, len(documents)